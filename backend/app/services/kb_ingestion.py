"""
Knowledge Base Ingestion Pipeline

Handles document parsing, chunking, embedding generation, and pgvector storage
for the Knowledge Base RAG system.

Design principles:
- Use lightweight parsers to avoid heavy dependencies
- Route embeddings through Bonito's own gateway (customer's models)
- Support all major document formats (PDF, DOCX, TXT, MD, HTML, CSV)
- Configurable chunking with overlap for context continuity
"""

import hashlib
import logging
import uuid
import asyncio
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional, Tuple, AsyncGenerator
from io import BytesIO
import mimetypes

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update

from app.models.knowledge_base import KnowledgeBase, KBDocument, KBChunk
from app.core.database import get_db_session

logger = logging.getLogger(__name__)


class DocumentParser:
    """Document parsing for various file types using lightweight libraries."""
    
    @staticmethod
    def parse_pdf(content: bytes) -> str:
        """Parse PDF content using PyPDF2/pypdf."""
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            text_parts = []
            
            for page in pdf_reader.pages:
                try:
                    text_parts.append(page.extract_text())
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page: {e}")
                    continue
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")
    
    @staticmethod
    def parse_docx(content: bytes) -> str:
        """Parse DOCX content using python-docx."""
        try:
            import docx
            from io import BytesIO
            
            doc = docx.Document(BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    @staticmethod
    def parse_html(content: bytes) -> str:
        """Parse HTML content using BeautifulSoup."""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text and clean up whitespace
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except ImportError:
            raise ImportError("beautifulsoup4 is required for HTML parsing. Install with: pip install beautifulsoup4")
        except Exception as e:
            raise ValueError(f"Failed to parse HTML: {e}")
    
    @staticmethod
    def parse_markdown(content: bytes) -> str:
        """Parse Markdown content."""
        try:
            import markdown
            from bs4 import BeautifulSoup
            
            # Convert markdown to HTML, then extract text
            html = markdown.markdown(content.decode('utf-8'))
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
            
        except ImportError:
            # Fallback: treat as plain text
            return content.decode('utf-8')
        except Exception as e:
            raise ValueError(f"Failed to parse Markdown: {e}")
    
    @staticmethod
    def parse_csv(content: bytes) -> str:
        """Parse CSV content."""
        try:
            import csv
            from io import StringIO
            
            text = content.decode('utf-8')
            csv_reader = csv.reader(StringIO(text))
            
            rows = []
            for row in csv_reader:
                rows.append(" | ".join(row))
            
            return "\n".join(rows)
            
        except Exception as e:
            raise ValueError(f"Failed to parse CSV: {e}")
    
    @staticmethod
    def parse_text(content: bytes) -> str:
        """Parse plain text content."""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file")
    
    @classmethod
    def parse_document(cls, content: bytes, file_type: str) -> str:
        """Parse document based on file type."""
        parsers = {
            'pdf': cls.parse_pdf,
            'docx': cls.parse_docx,
            'html': cls.parse_html,
            'htm': cls.parse_html,
            'md': cls.parse_markdown,
            'markdown': cls.parse_markdown,
            'csv': cls.parse_csv,
            'txt': cls.parse_text,
            'json': cls.parse_text,  # JSON is treated as text
        }
        
        parser = parsers.get(file_type.lower())
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return parser(content)


class TextChunker:
    """Text chunking with configurable size and overlap."""
    
    def __init__(self, chunk_size: int = 512, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def estimate_tokens(self, text: str) -> int:
        """Rough token estimation: ~4 characters per token for English."""
        return len(text) // 4
    
    def split_text_recursive(self, text: str, separators: List[str] = None) -> List[str]:
        """
        Recursively split text using different separators.
        
        Order of preference:
        1. Paragraphs (\n\n)
        2. Lines (\n)
        3. Sentences (. )
        4. Words ( )
        5. Characters
        """
        if separators is None:
            separators = ["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        
        if not text.strip():
            return []
        
        # Check if text is already small enough
        if self.estimate_tokens(text) <= self.chunk_size:
            return [text.strip()]
        
        if not separators:
            # No more separators, split by characters
            chunks = []
            for i in range(0, len(text), self.chunk_size * 4):  # 4 chars per token
                chunk = text[i:i + self.chunk_size * 4]
                if chunk.strip():
                    chunks.append(chunk.strip())
            return chunks
        
        separator = separators[0]
        remaining_separators = separators[1:]
        
        # Split by current separator
        splits = text.split(separator)
        
        # Recombine splits into chunks
        chunks = []
        current_chunk = ""
        
        for split in splits:
            if not split.strip():
                continue
            
            # Would adding this split exceed the chunk size?
            potential_chunk = current_chunk + (separator if current_chunk else "") + split
            
            if self.estimate_tokens(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
            else:
                # Current chunk is full, start a new one
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # If the split itself is too large, recursively split it
                if self.estimate_tokens(split) > self.chunk_size:
                    sub_chunks = self.split_text_recursive(split, remaining_separators)
                    chunks.extend(sub_chunks)
                    current_chunk = ""
                else:
                    current_chunk = split.strip()
        
        # Don't forget the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def add_overlap(self, chunks: List[str]) -> List[str]:
        """Add overlap between consecutive chunks."""
        if len(chunks) <= 1 or self.overlap <= 0:
            return chunks
        
        overlapped_chunks = []
        
        for i, chunk in enumerate(chunks):
            if i == 0:
                overlapped_chunks.append(chunk)
                continue
            
            # Get overlap from previous chunk
            prev_chunk = chunks[i-1]
            prev_words = prev_chunk.split()
            
            if len(prev_words) >= self.overlap:
                overlap_text = " ".join(prev_words[-self.overlap:])
                overlapped_chunk = overlap_text + " " + chunk
            else:
                overlapped_chunk = prev_chunk + " " + chunk
            
            overlapped_chunks.append(overlapped_chunk)
        
        return overlapped_chunks
    
    def chunk_text(self, text: str) -> List[str]:
        """Main chunking method."""
        if not text.strip():
            return []
        
        # Split recursively
        chunks = self.split_text_recursive(text)
        
        # Add overlap
        chunks = self.add_overlap(chunks)
        
        # Filter out empty chunks
        return [chunk for chunk in chunks if chunk.strip()]


class EmbeddingGenerator:
    """Generate embeddings through Bonito's gateway using customer's models.
    
    Automatically selects the cheapest AVAILABLE embedding model from the
    org's connected providers. Only considers models that are actually
    routable (have credentials + are serverless or explicitly deployed).
    Falls back through multiple models with a timeout to avoid hanging.
    """
    
    EMBEDDING_TIMEOUT = 30  # seconds per batch — fail fast, don't hang
    
    def __init__(self, org_id: uuid.UUID):
        self.org_id = org_id
    
    async def get_cheapest_embedding_model(self, db: AsyncSession) -> str:
        """
        Determine the cheapest available embedding model for the organization.
        
        Only picks models that are actually routable through the gateway.
        Serverless models (GCP Vertex, OpenAI) are preferred because they
        don't require explicit deployment. AWS Bedrock models need activation.
        """
        from app.services.gateway import get_router
        
        # Priority order: serverless-first (GCP/OpenAI always available if API enabled),
        # then AWS (may need model activation), then others
        preferred_models = [
            # GCP — serverless, always available if Vertex AI API is enabled
            "text-embedding-005",
            "text-multilingual-embedding-002",
            "gemini-embedding-001",
            # OpenAI — serverless, always available with API key
            "text-embedding-3-small",
            # AWS — may need model access activation first
            "amazon.titan-embed-text-v2:0",
            "amazon.titan-embed-text-v1",
            "cohere.embed-english-v3",
            "cohere.embed-v4:0",
        ]
        
        try:
            router = await get_router(db, self.org_id)
            available = {m["model_name"] for m in router.model_list} if router.model_list else set()
            
            for model in preferred_models:
                if model in available:
                    logger.info(f"Selected embedding model: {model}")
                    return model
            
            # If none of our preferred models match, try ANY embedding model in the router
            for m in router.model_list or []:
                name = m.get("model_name", "").lower()
                if "embed" in name or "embedding" in name:
                    logger.info(f"Selected fallback embedding model: {m['model_name']}")
                    return m["model_name"]
                    
        except Exception as e:
            logger.warning(f"Failed to query available models: {e}")
        
        # Last resort — GCP is most likely to work serverlessly
        return "text-embedding-005"
    
    async def generate_embeddings(self, texts: List[str], model: str = None, dimensions: int = None) -> List[List[float]]:
        """
        Generate embeddings for a list of texts.
        
        Routes through the org's LiteLLM router to use their configured
        embedding models and cloud credentials. Includes timeout to prevent
        hanging on unresponsive models (e.g., unactivated Bedrock models).
        
        Args:
            dimensions: If set, request reduced-dimension embeddings (supported by
                        OpenAI text-embedding-3-small/large). Useful when the KB
                        vector column is smaller than the model's default output.
        """
        if not texts:
            return []
        
        if model is None:
            async with get_db_session() as db:
                model = await self.get_cheapest_embedding_model(db)
        
        logger.info(f"Generating embeddings for {len(texts)} texts using model {model} (dims={dimensions})")
        
        from app.services.gateway import get_router
        
        embeddings = []
        batch_size = 20  # Process in batches to avoid rate limits
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            try:
                async with get_db_session() as db:
                    router = await get_router(db, self.org_id)
                    # Timeout to prevent hanging on unactivated/unavailable models
                    embed_kwargs = {"model": model, "input": batch}
                    if dimensions:
                        embed_kwargs["dimensions"] = dimensions
                    response = await asyncio.wait_for(
                        router.aembedding(**embed_kwargs),
                        timeout=self.EMBEDDING_TIMEOUT,
                    )
                    for item in response.data:
                        embeddings.append(item["embedding"])
            except asyncio.TimeoutError:
                logger.error(f"Embedding timed out after {self.EMBEDDING_TIMEOUT}s for model {model} (batch {i//batch_size}). "
                             f"Model may not be activated — check provider console or use one-click activation.")
                raise RuntimeError(
                    f"Embedding model '{model}' timed out. It may not be activated on your cloud provider. "
                    f"Use Bonito's model activation feature or choose a different embedding model."
                )
            except Exception as e:
                logger.error(f"Embedding generation failed for batch {i//batch_size}: {e}")
                raise
            
            # Small delay between batches to avoid rate limits
            if i + batch_size < len(texts):
                await asyncio.sleep(0.5)
        
        logger.info(f"Generated {len(embeddings)} embeddings successfully")
        return embeddings


async def process_document(
    doc_id: uuid.UUID,
    content: bytes,
    kb_id: uuid.UUID,
    force_reprocess: bool = False
) -> None:
    """
    Process a document through the full ingestion pipeline.
    
    Steps:
    1. Parse document content
    2. Split into chunks
    3. Generate embeddings
    4. Store in database
    """
    async with get_db_session() as db:
        # Get document and knowledge base info
        doc_result = await db.execute(
            select(KBDocument).where(KBDocument.id == doc_id)
        )
        doc = doc_result.scalar_one_or_none()
        if not doc:
            logger.error(f"Document {doc_id} not found")
            return
        
        kb_result = await db.execute(
            select(KnowledgeBase).where(KnowledgeBase.id == kb_id)
        )
        kb = kb_result.scalar_one_or_none()
        if not kb:
            logger.error(f"Knowledge base {kb_id} not found")
            return
        
        try:
            # Mark as processing
            doc.status = "processing"
            await db.commit()
            
            logger.info(f"Processing document {doc_id} ({doc.file_name})")
            
            # Step 1: Parse document content
            parsed_text = DocumentParser.parse_document(content, doc.file_type or "txt")
            
            if not parsed_text.strip():
                raise ValueError("No text content found in document")
            
            # Step 2: Split into chunks
            chunker = TextChunker(
                chunk_size=kb.chunk_size,
                overlap=kb.chunk_overlap
            )
            chunks = chunker.chunk_text(parsed_text)
            
            if not chunks:
                raise ValueError("No chunks generated from document")
            
            # Step 3: Generate embeddings
            embedding_gen = EmbeddingGenerator(doc.org_id)
            embed_model = kb.embedding_model if kb.embedding_model != "auto" else None
            # Pass target dimensions to support OpenAI dimension reduction
            # (e.g., text-embedding-3-small can produce 768 instead of default 1536)
            embed_dims = kb.embedding_dimensions if kb.embedding_dimensions else None
            embeddings = await embedding_gen.generate_embeddings(
                chunks, 
                model=embed_model,
                dimensions=embed_dims
            )
            
            if len(embeddings) != len(chunks):
                raise ValueError(f"Embedding count ({len(embeddings)}) doesn't match chunk count ({len(chunks)})")
            
            # Step 4: Store chunks in database
            chunk_objects = []
            total_tokens = 0
            
            for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                token_count = chunker.estimate_tokens(chunk_text)
                total_tokens += token_count
                
                chunk_obj = KBChunk(
                    document_id=doc_id,
                    knowledge_base_id=kb_id,
                    org_id=doc.org_id,
                    content=chunk_text,
                    token_count=token_count,
                    chunk_index=i,
                    embedding=embedding,
                    source_file=doc.file_name,
                    source_page=None,  # TODO: Extract page numbers during parsing
                    source_section=None,  # TODO: Extract sections during parsing
                    extra_metadata={"processed_at": datetime.now(timezone.utc).isoformat()}
                )
                chunk_objects.append(chunk_obj)
            
            # Bulk insert chunks
            db.add_all(chunk_objects)
            
            # Update document status
            doc.status = "ready"
            doc.chunk_count = len(chunks)
            doc.updated_at = datetime.now(timezone.utc)
            
            # Update knowledge base counters
            kb.document_count += 1
            kb.chunk_count += len(chunks)
            kb.total_tokens += total_tokens
            if kb.status == "pending":
                kb.status = "ready"
            
            await db.commit()
            
            logger.info(f"Successfully processed document {doc_id}: {len(chunks)} chunks, {total_tokens} tokens")
            
        except Exception as e:
            logger.error(f"Failed to process document {doc_id}: {e}")
            
            # Mark as error
            doc.status = "error"
            doc.error_message = str(e)[:1000]  # Truncate long error messages
            doc.updated_at = datetime.now(timezone.utc)
            
            # Update KB status if this was the only document
            if kb.document_count == 1:
                kb.status = "error"
            
            await db.commit()
            raise


async def sync_from_cloud_storage(kb_id: uuid.UUID) -> None:
    """
    Sync documents from cloud storage (S3, Azure Blob, GCS).
    Delegates to storage_connector.sync_kb_from_storage().
    """
    from app.services.storage_connector import sync_kb_from_storage
    await sync_kb_from_storage(kb_id)
    await asyncio.sleep(1)  # Simulate work


async def search_chunks(
    kb_id: uuid.UUID,
    query_embedding: List[float],
    top_k: int = 5,
    min_score: float = 0.0,
    org_id: uuid.UUID = None
) -> List[Dict[str, Any]]:
    """
    Search for similar chunks using pgvector cosine similarity.
    
    TODO: Implement actual vector search
    """
    async with get_db_session() as db:
        # Verify knowledge base access
        if org_id:
            kb_result = await db.execute(
                select(KnowledgeBase).where(
                    KnowledgeBase.id == kb_id,
                    KnowledgeBase.org_id == org_id
                )
            )
            kb = kb_result.scalar_one_or_none()
            if not kb:
                raise ValueError("Knowledge base not found or access denied")
        
        # TODO: Implement pgvector similarity search
        # Example query:
        # SELECT *, (embedding <=> %s) as distance 
        # FROM kb_chunks 
        # WHERE knowledge_base_id = %s 
        # ORDER BY distance ASC 
        # LIMIT %s
        
        # Placeholder: return empty results
        logger.info(f"Vector search in KB {kb_id} for {len(query_embedding)}-dim embedding (top_k={top_k})")
        return []